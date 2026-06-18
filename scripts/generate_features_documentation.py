from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Teams_SharePoint_AI_Assistant_Features_Documentation.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    run = title.add_run("Teams SharePoint AI Assistant")
    run.bold = True
    run.font.size = Pt(22)
    subtitle = doc.add_paragraph()
    subtitle.add_run("Features Documentation").bold = True
    doc.add_paragraph(f"Generated: {date.today().isoformat()}")

    add_heading(doc, "1. Overview", 1)
    doc.add_paragraph(
        "The Teams SharePoint AI Assistant is a Microsoft Teams bot that helps users "
        "search, summarize, review, and work with organizational SharePoint documents. "
        "It is designed to be SharePoint-first, cache-aware, source-grounded, and capable "
        "of handling natural follow-up questions."
    )

    features = [
        (
            "Microsoft Teams Chat Assistant",
            [
                "Runs inside Microsoft Teams as a conversational bot.",
                "Supports one-on-one and group/channel conversations.",
                "Sends typing indicators quickly while processing requests.",
                "Provides concise workplace-style responses.",
            ],
        ),
        (
            "SharePoint-First Knowledge Retrieval",
            [
                "Uses configured SharePoint sites as the main organizational knowledge source.",
                "Searches cached SharePoint documents first for fast answers.",
                "Uses live Microsoft Graph SharePoint search/listing as fallback when needed.",
                "Keeps OneDrive, Azure AI Search, and web search disabled by default unless enabled through environment variables.",
            ],
        ),
        (
            "Local SharePoint Document Cache",
            [
                "Stores extracted document text and metadata locally.",
                "Preserves document title, URL, content, source metadata, and visibility.",
                "Speeds up repeated document questions and summaries.",
                "Supports document hydration before prompt building.",
            ],
        ),
        (
            "Document Search And Ranking",
            [
                "Prioritizes exact and strong title matches.",
                "Ignores helper words such as can, you, please, summarize, document, and file during relevance scoring.",
                "Avoids unrelated documents when a strong title match exists.",
            ],
        ),
        (
            "Full Document Summary Mode",
            [
                "Detects summary and overview requests.",
                "Uses larger bounded context windows for exact or strong title matches.",
                "Prefers cached full document content over snippets.",
                "Adds truncation notes when only part of a document is included.",
            ],
        ),
        (
            "Follow-Up Question Handling",
            [
                "Tracks previous source documents and recent conversation history.",
                "Routes clear follow-ups to previous context instead of starting a new search.",
                "Supports follow-ups such as 'tell me more about it', 'I meant the employee handbook', and 'what should I add'.",
            ],
        ),
        (
            "Document Review And Recommendations",
            [
                "Supports document-improvement questions.",
                "Suggests additions, gaps, and improvements based on included document content.",
                "Cites the source document when describing what it already contains.",
                "Labels recommendations clearly as recommendations.",
            ],
        ),
        (
            "Document Title Listing And Pagination",
            [
                "Lists available cached SharePoint document titles directly from metadata.",
                "Supports requests such as 'list top 10 document titles' and 'show available documents'.",
                "Supports pagination for large libraries using 'show more' or 'next 10'.",
            ],
        ),
        (
            "Listed Document Summary With URLs",
            [
                "Remembers the last shown title batch.",
                "Summarizes those exact listed titles when the user asks for short summaries and URLs.",
                "Uses cached document content to create compact summaries.",
                "Includes available document URLs.",
            ],
        ),
        (
            "Uploaded File Support",
            [
                "Processes uploaded Teams attachments.",
                "Supports common document and data formats.",
                "Stores uploaded content for follow-up questions.",
                "Prioritizes uploaded files over external search when the user asks about an uploaded file.",
            ],
        ),
        (
            "Source Citations",
            [
                "Requires citations for organizational facts.",
                "Avoids citing documents that were not actually used.",
                "Avoids inventing organizational facts when source content is unavailable.",
            ],
        ),
        (
            "Context And Token Budget Controls",
            [
                "Compresses and caps document content before sending it to the LLM.",
                "Applies per-document and total-context limits.",
                "Uses larger limits for document summaries where appropriate.",
                "Applies a final token gatekeeper before model calls.",
            ],
        ),
        (
            "Logging And Diagnostics",
            [
                "Logs routing decisions, search scope, selected primary document, content sizes, and truncation status.",
                "Logs cache behavior and Graph/listing failures for operations support.",
            ],
        ),
    ]

    add_heading(doc, "2. Feature Details", 1)
    for idx, (name, bullets) in enumerate(features, 1):
        add_heading(doc, f"{idx}. {name}", 2)
        add_bullets(doc, bullets)

    add_heading(doc, "3. Feature Configuration Summary", 1)
    add_table(
        doc,
        ["Feature", "Main Setting"],
        [
            ["SharePoint search", "ENABLE_SHAREPOINT_SEARCH=true"],
            ["SharePoint cache", "ENABLE_SHAREPOINT_CACHE=true"],
            ["Cache-first behavior", "SHAREPOINT_CACHE_FIRST=true"],
            ["SharePoint-only mode", "DATA_SOURCE_MODE=sharepoint"],
            ["OneDrive search", "ENABLE_ONEDRIVE_SEARCH=false by default"],
            ["Azure AI Search", "ENABLE_AI_SEARCH=false by default"],
            ["Web search", "ENABLE_WEB_SEARCH=false by default"],
            ["Startup SharePoint crawl", "ENABLE_SHAREPOINT_STARTUP_CRAWL=true"],
            ["Summary context windows", "SUMMARY_PRIMARY_DOC_CHARS, SUMMARY_DOC_SNIPPET_CHARS, SUMMARY_TOTAL_CONTEXT_CHARS"],
            ["Prompt safety limits", "MAX_PROMPT_CHARS, MAX_DOCS, MAX_DOC_SNIPPET_CHARS, MAX_TOTAL_CONTEXT_CHARS"],
        ],
    )

    add_heading(doc, "4. Packaging Note", 1)
    doc.add_paragraph(
        "This features documentation is part of the project handover/deployment documentation package. "
        "It should not be inserted into the Microsoft Teams app package zip. The Teams app package "
        "should contain the Teams manifest and required icon assets only. Documentation should be "
        "distributed alongside the deployment package, stored in the project repository, or uploaded "
        "to the customer handover SharePoint location."
    )

    add_heading(doc, "5. Feature Acceptance Criteria", 1)
    add_bullets(
        doc,
        [
            "Exact document-title questions return the intended document first.",
            "Summary requests include enough document content to be useful.",
            "Follow-up questions do not trigger unnecessary new searches.",
            "Title listing returns real cached SharePoint titles.",
            "Title listing pagination works with show more.",
            "Listed-title summary requests summarize the last listed batch.",
            "Organizational facts are cited.",
            "Optional data sources remain disabled unless explicitly enabled.",
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
