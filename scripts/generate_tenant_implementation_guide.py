from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Teams_SharePoint_AI_Assistant_Tenant_Implementation_Guide.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def build_doc() -> None:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    run = title.add_run("Teams SharePoint AI Assistant")
    run.bold = True
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.add_run("Tenant Implementation Guide and Requirements").bold = True
    doc.add_paragraph(f"Generated: {date.today().isoformat()}")
    doc.add_paragraph(
        "This document explains what the project does and provides a step-by-step "
        "process for implementing it in another Microsoft 365/Azure tenant."
    )

    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "The Teams SharePoint AI Assistant is a Microsoft Teams bot that answers "
        "questions from organizational SharePoint documents. It uses Azure OpenAI "
        "for routing, summarization, and response generation, Microsoft Graph for "
        "SharePoint access, and a local document cache for fast document retrieval."
    )
    add_bullets(
        doc,
        [
            "Primary source: configured SharePoint libraries and cached SharePoint documents.",
            "Default behavior: SharePoint-only, with OneDrive, web, and Azure AI Search disabled unless explicitly enabled.",
            "Main value: fast, cited answers from organizational documents inside Microsoft Teams.",
            "Key safeguards: source citations, follow-up context reuse, title-match ranking, and prompt-size controls.",
        ],
    )

    add_heading(doc, "2. What The Project Does", 1)
    add_bullets(
        doc,
        [
            "Runs as a Microsoft Teams bot.",
            "Uses Azure OpenAI to understand user requests and generate answers.",
            "Searches cached SharePoint content first for faster responses.",
            "Falls back to live Microsoft Graph SharePoint search/listing when needed.",
            "Summarizes Word, PDF, PowerPoint, Excel, CSV, text, JSON, and XML files when extractable.",
            "Supports uploaded files in Teams conversations.",
            "Tracks prior source documents so follow-up questions do not unnecessarily search again.",
            "Lists available document titles and supports pagination for large SharePoint libraries.",
            "Provides citations for organizational facts.",
        ],
    )

    add_heading(doc, "3. Target Architecture", 1)
    add_table(
        doc,
        ["Component", "Purpose"],
        [
            ["Microsoft Teams app/bot", "User-facing chat interface."],
            ["Azure Bot registration", "Bot identity, Teams channel connection, and messaging endpoint registration."],
            ["Azure App Service", "Hosts the Python Teams bot application."],
            ["Azure OpenAI", "Provides the deployed chat model used for routing and answer generation."],
            ["Microsoft Graph app registration", "Provides app-only access to configured SharePoint sites/libraries."],
            ["SharePoint site/libraries", "Primary organizational knowledge source."],
            ["Local document cache", "Stores extracted SharePoint text for fast retrieval and follow-up answers."],
            ["Application settings/environment variables", "Stores non-code configuration and secrets in the target tenant."],
        ],
    )

    add_heading(doc, "4. Access And Ownership Requirements", 1)
    add_table(
        doc,
        ["Requirement", "Needed Access"],
        [
            ["Azure subscription", "Contributor or Owner on the target resource group, or equivalent deployment rights."],
            ["Azure App Service", "Permission to create/update App Service, configure app settings, and view logs."],
            ["Azure OpenAI", "Access to an Azure OpenAI resource and a deployed model."],
            ["Microsoft Entra ID", "Permission to create app registrations or work with an Entra admin."],
            ["Microsoft Graph permissions", "Admin consent for required Graph application permissions."],
            ["SharePoint", "Access to target SharePoint sites and permission to grant the app access."],
            ["Teams admin", "Permission to upload/publish the Teams app package or approve app installation."],
            ["Service account or managed owner", "A tenant-owned operational identity for administration and support."],
        ],
    )

    add_heading(doc, "5. Step-By-Step Tenant Implementation Process", 1)

    add_heading(doc, "Step 1 - Confirm Scope And Source Libraries", 2)
    add_numbered(
        doc,
        [
            "Identify the target tenant and business owner.",
            "List SharePoint sites and document libraries the assistant should use.",
            "Decide whether the bot should be available to a pilot group, one Team, or the full organization.",
            "Confirm which optional sources should stay disabled. Recommended initial setup is SharePoint-only.",
        ],
    )

    add_heading(doc, "Step 2 - Prepare Azure Resources", 2)
    add_numbered(
        doc,
        [
            "Create or select an Azure resource group in the target tenant/subscription.",
            "Create an Azure App Service Plan suitable for the expected traffic.",
            "Create an Azure App Service for the Python bot.",
            "Enable application logging and configure log retention.",
            "Confirm outbound network access to Microsoft Graph and Azure OpenAI endpoints.",
        ],
    )

    add_heading(doc, "Step 3 - Prepare Azure OpenAI", 2)
    add_numbered(
        doc,
        [
            "Create or select an Azure OpenAI resource.",
            "Deploy the required chat model.",
            "Record the Azure OpenAI endpoint, deployment name, and API key.",
            "Confirm quota and rate limits are sufficient for Teams usage.",
            "Set conservative token and concurrency settings for the first rollout.",
        ],
    )

    add_heading(doc, "Step 4 - Create The Bot App Registration", 2)
    add_numbered(
        doc,
        [
            "Create a Microsoft Entra app registration for the Teams bot identity.",
            "Generate a client secret and store it securely.",
            "Create or update the Azure Bot registration using the bot app ID.",
            "Enable the Microsoft Teams channel on the bot registration.",
            "Set the messaging endpoint to the deployed App Service endpoint, for example https://<app-name>.azurewebsites.net/api/messages if that is the configured route.",
        ],
    )

    add_heading(doc, "Step 5 - Create The Graph Access App Registration", 2)
    add_numbered(
        doc,
        [
            "Create a dedicated Microsoft Entra app registration for Microsoft Graph access, or confirm the bot app registration will also be used for Graph.",
            "Add the required Microsoft Graph application permissions.",
            "Recommended least-privilege approach: use Sites.Selected where possible and grant access only to required SharePoint sites.",
            "Alternative broad approach: Sites.Read.All or Files.Read.All, only if approved by tenant security policy.",
            "Have an administrator grant admin consent.",
            "Generate and securely store the Graph client secret.",
        ],
    )

    add_heading(doc, "Step 6 - Grant SharePoint Site Access", 2)
    add_numbered(
        doc,
        [
            "If using Sites.Selected, grant the Graph app read access to each required SharePoint site.",
            "Validate that the app can list files from each target library.",
            "Confirm the app can download supported document types for extraction.",
            "Document the approved site URLs for the SHAREPOINT_SITES setting.",
        ],
    )

    add_heading(doc, "Step 7 - Configure App Service Settings", 2)
    add_table(
        doc,
        ["Setting", "Purpose"],
        [
            ["BOT_ID", "Teams bot app/client ID."],
            ["SECRET_BOT_PASSWORD or BOT_PASSWORD", "Teams bot client secret."],
            ["TENANT_ID or TEAMS_APP_TENANT_ID", "Target tenant ID."],
            ["GRAPH_CLIENT_ID", "Graph app registration client ID."],
            ["GRAPH_CLIENT_SECRET", "Graph app client secret."],
            ["GRAPH_TENANT_ID", "Tenant ID used for Graph token acquisition."],
            ["AZURE_OPENAI_ENDPOINT", "Azure OpenAI endpoint."],
            ["AZURE_OPENAI_API_KEY", "Azure OpenAI API key."],
            ["AZURE_OPENAI_MODEL_DEPLOYMENT_NAME", "Azure OpenAI model deployment name."],
            ["SHAREPOINT_SITES", "Comma-separated SharePoint site URLs."],
            ["DATA_SOURCE_MODE", "Recommended: sharepoint."],
            ["ENABLE_SHAREPOINT_SEARCH", "Recommended: true."],
            ["ENABLE_SHAREPOINT_CACHE", "Recommended: true."],
            ["SHAREPOINT_CACHE_FIRST", "Recommended: true."],
            ["ENABLE_ONEDRIVE_SEARCH", "Recommended initial value: false."],
            ["ENABLE_AI_SEARCH", "Recommended initial value: false."],
            ["ENABLE_WEB_SEARCH", "Recommended initial value: false."],
        ],
    )

    add_heading(doc, "Step 8 - Deploy The Application", 2)
    add_numbered(
        doc,
        [
            "Deploy the application code to Azure App Service.",
            "Install dependencies from requirements.txt.",
            "Configure the startup command for the Python Teams bot.",
            "Restart the App Service after application settings are saved.",
            "Open App Service logs and confirm startup completes without missing environment variable errors.",
        ],
    )

    add_heading(doc, "Step 9 - Create Or Update The Teams App Package", 2)
    add_numbered(
        doc,
        [
            "Update the Teams app manifest with the target bot ID, app ID, valid domains, and bot endpoint domain.",
            "Package the manifest and icons into the Teams app package.",
            "Upload the package to Teams for testing or submit it to the tenant app catalog.",
            "Install the app for pilot users or the target Team.",
        ],
    )

    add_heading(doc, "Step 10 - Warm The SharePoint Cache", 2)
    add_numbered(
        doc,
        [
            "Start the app with SharePoint cache enabled.",
            "If startup crawling is enabled, let the app index the configured SharePoint libraries.",
            "If startup crawling is disabled, cache will populate as documents are used.",
            "Validate that document_cache.json or the configured cache store contains expected document metadata and content.",
            "Run test prompts for known documents such as employee handbooks, policies, guides, or implementation documents.",
        ],
    )

    add_heading(doc, "Step 11 - Validate Core Scenarios", 2)
    add_bullets(
        doc,
        [
            "Ask for a known document by exact title and verify it ranks first.",
            "Ask for a full summary and verify the answer uses enough document content.",
            "Ask a follow-up question and verify no new SharePoint search is triggered.",
            "Ask for top 10 document titles and verify metadata listing works.",
            "Ask show more and verify pagination works.",
            "Ask for short summaries and URLs for listed titles and verify the same listed batch is used.",
            "Upload a supported file and ask a question about it.",
            "Check citations, truncation notes, and logs.",
        ],
    )

    add_heading(doc, "Step 12 - Pilot, Monitor, And Handover", 2)
    add_numbered(
        doc,
        [
            "Run a pilot with a small user group.",
            "Monitor App Service logs, Azure OpenAI usage, Teams response latency, and Graph failures.",
            "Tune context limits, result counts, and crawl settings based on real usage.",
            "Create a support runbook for secret rotation, app restarts, cache refresh, and SharePoint permission changes.",
            "Hand over ownership to the tenant operations team.",
        ],
    )

    add_heading(doc, "6. Recommended Initial Environment Values", 1)
    add_table(
        doc,
        ["Variable", "Recommended Value"],
        [
            ["DATA_SOURCE_MODE", "sharepoint"],
            ["ENABLE_SHAREPOINT_SEARCH", "true"],
            ["ENABLE_SHAREPOINT_CACHE", "true"],
            ["ENABLE_SHAREPOINT_STARTUP_CRAWL", "true for small/medium libraries; false for controlled/manual cache warmup"],
            ["ENABLE_LIVE_GRAPH_FALLBACK", "true"],
            ["ENABLE_ONEDRIVE_SEARCH", "false"],
            ["ENABLE_AI_SEARCH", "false"],
            ["ENABLE_WEB_SEARCH", "false"],
            ["SHAREPOINT_CACHE_FIRST", "true"],
            ["MAX_DOCS", "4 to 6"],
            ["MAX_PROMPT_CHARS", "60000 to 120000 depending on model and latency target"],
            ["LLM_CONCURRENCY", "1 for early rollout"],
        ],
    )

    add_heading(doc, "7. Security And Compliance Notes", 1)
    add_bullets(
        doc,
        [
            "Do not place secrets in source control or documentation.",
            "Use Azure App Service Application Settings or Key Vault references for secrets.",
            "Prefer least-privilege SharePoint access using Sites.Selected when possible.",
            "Do not enable OneDrive, AI Search, or web search unless approved.",
            "Review document cache handling with the tenant security team.",
            "Confirm retention expectations for cached extracted document text.",
            "Rotate bot and Graph secrets according to tenant policy.",
        ],
    )

    add_heading(doc, "8. Testing Checklist", 1)
    add_bullets(
        doc,
        [
            "Bot starts without missing configuration errors.",
            "Teams can send and receive messages from the bot.",
            "Typing indicator appears quickly.",
            "Known SharePoint document search returns the expected title.",
            "Document summary includes citations.",
            "Follow-up questions reuse previous context and do not search again.",
            "Document title listing returns more than one title when cache contains many documents.",
            "Pagination works with show more.",
            "Listed-title summary uses the same titles that were just listed.",
            "Logs show selected primary document, content passed to LLM, and truncation status.",
        ],
    )

    add_heading(doc, "9. Risks And Mitigations", 1)
    add_table(
        doc,
        ["Risk", "Mitigation"],
        [
            ["Graph permissions too broad", "Use Sites.Selected or restrict SharePoint sites where possible."],
            ["Slow responses", "Use cache-first retrieval, reduce result counts, and tune prompt limits."],
            ["Wrong documents selected", "Use title-first ranking and strong title-match filtering."],
            ["Follow-ups trigger new searches", "Use previous-source follow-up routing and final no-search guard."],
            ["Secrets exposed", "Store secrets in App Service settings or Key Vault; never include values in docs."],
            ["Cache contains stale content", "Define a cache refresh process and monitor crawl/update timestamps."],
        ],
    )

    add_heading(doc, "10. Acceptance Criteria", 1)
    add_bullets(
        doc,
        [
            "The bot is installable and usable in Microsoft Teams.",
            "The bot answers from the configured SharePoint libraries.",
            "The bot can summarize a known document by exact title.",
            "The bot can list document titles and paginate through large libraries.",
            "Follow-up questions use previous context without unnecessary search.",
            "Organizational facts are cited.",
            "The deployment uses tenant-owned Azure resources and tenant-approved permissions.",
            "Operational owners can rotate secrets, update SharePoint sites, restart the app, and review logs.",
        ],
    )

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "Appendix - Implementation Inputs To Collect", 1)
    add_table(
        doc,
        ["Input", "Value To Collect"],
        [
            ["Tenant ID", ""],
            ["Azure subscription ID", ""],
            ["Resource group", ""],
            ["App Service name", ""],
            ["Bot app registration ID", ""],
            ["Graph app registration ID", ""],
            ["Azure OpenAI resource name", ""],
            ["Azure OpenAI endpoint", ""],
            ["Azure OpenAI deployment name", ""],
            ["SharePoint site URLs", ""],
            ["Teams app package owner", ""],
            ["Operations owner", ""],
            ["Security approver", ""],
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
